import fitz  # PyMuPDF
from docx import Document
from pathlib import Path
from typing import Dict, List, Set, Tuple

# Images smaller than this in both dimensions are treated as decorative
# icons (checkbox/arrow/logo glyphs) and skipped entirely.
MIN_IMAGE_DIM = 100


def _pdf_text_lines(page) -> List[Tuple[float, float, str]]:
    """
    Rebuild the page's text line-by-line from get_text("words").
    Grouping by (block_no, line_no) and joining with spaces reproduces
    the flat get_text("text") output exactly while also giving each line
    a y/x position for ordering against images.
    """
    words = page.get_text("words")  # (x0, y0, x1, y1, word, block_no, line_no, word_no)

    grouped = {}
    for w in words:
        key = (w[5], w[6])
        grouped.setdefault(key, []).append(w)

    lines = []
    for key in sorted(grouped):
        ws = sorted(grouped[key], key=lambda p: p[0])
        lines.append((
            min(w[1] for w in ws),   # y0
            min(w[0] for w in ws),   # x0
            " ".join(w[4] for w in ws),
        ))
    return lines


def _pdf_rendered_images(page) -> List[Tuple[int, float, float]]:
    """
    Return (xref, y0, x0) for each distinct image actually rendered on a page.
    get_image_info(xrefs=True) reports only images placed on the page, unlike
    get_images(full=True) which can list shared xrefs on every page.
    """
    seen = set()
    items = []
    for info in page.get_image_info(xrefs=True):
        xref = info.get("xref")
        if not xref or xref in seen:
            continue
        seen.add(xref)
        bbox = info.get("bbox") or (0.0, 0.0, 0.0, 0.0)
        items.append((xref, bbox[1], bbox[0]))
    return items


def parse_pdf(file_path: str, output_image_dir: str) -> Dict:
    """
    Extract text and images from PDF, preserving per-page structure.
    Image tags like ![Image](page_X_img_Y.png) are embedded in the text
    at the position where the image appears on the page.

    Filtering rules:
      - Only images actually rendered on a page are considered
        (via page.get_image_info, avoiding shared xrefs listed on every page).
      - Decorative images are skipped entirely: anything smaller than
        MIN_IMAGE_DIM in both dimensions, or whose xref renders on more than
        one page (repeated logos / watermarks).
      - Each image is stored once on disk and tagged inline next to the text
        it belongs to (ordered by position, not appended at the end).
    """
    doc = fitz.open(file_path)

    full_text = []
    image_paths = []
    pages = []
    pages_processed = doc.page_count

    output_dir = Path(output_image_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    # Pre-pass: count how many pages each xref renders on.
    xref_page_count = {}
    for page_index in range(pages_processed):
        page = doc[page_index]
        for xref, _, _ in _pdf_rendered_images(page):
            xref_page_count[xref] = xref_page_count.get(xref, 0) + 1

    xref_to_tag = {}
    written_xrefs: Set[int] = set()

    for page_index in range(pages_processed):
        page = doc[page_index]

        lines = _pdf_text_lines(page)

        # Keep only content images on this page.
        kept_images = []
        for xref, y0, x0 in _pdf_rendered_images(page):
            if xref_page_count.get(xref, 0) > 1:
                continue
            base_image = doc.extract_image(xref)
            if base_image["width"] < MIN_IMAGE_DIM and base_image["height"] < MIN_IMAGE_DIM:
                continue
            kept_images.append((xref, y0, x0))

        image_tags = []
        for img_index, (xref, y0, x0) in enumerate(kept_images):
            tag = xref_to_tag.get(xref)
            if tag is None:
                tag = f"page_{page_index}_img_{img_index}.png"
                base_image = doc.extract_image(xref)
                image_path = output_dir / tag
                with open(image_path, "wb") as f:
                    f.write(base_image["image"])
                image_paths.append(str(image_path))
                written_xrefs.add(xref)
                xref_to_tag[xref] = tag
            image_tags.append((y0, x0, tag))

        # Merge text lines and image tags by page position.
        items = [("text", y0, x0, text) for y0, x0, text in lines]
        items += [("image", y0, x0, f"![Image]({tag})") for y0, x0, tag in image_tags]
        items.sort(key=lambda it: (it[1], it[2]))

        text = "\n".join(item[3] for item in items)

        full_text.append(text)

        pages.append({
            "page_number": page_index + 1,
            "text": text,
        })

    return {
        "text": "\n".join(full_text),
        "pages": pages,
        "image_paths": image_paths,
        "pages_processed": pages_processed,
    }


def parse_docx(file_path: str, output_image_dir: str) -> Dict:
    """
    Extract text and images from DOCX.
    Image tags are embedded at the end of the text.
    """
    doc = Document(file_path)

    full_text = []
    image_paths = []
    image_tags = []

    output_dir = Path(output_image_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())

    text_content = "\n".join(full_text)

    rels = doc.part._rels
    for rel in rels:
        target = rels[rel].target_ref
        if "media" in target:
            image_data = rels[rel].target_part.blob
            tag = f"{rel}.png"
            image_path = output_dir / tag
            with open(image_path, "wb") as f:
                f.write(image_data)
            image_paths.append(str(image_path))
            image_tags.append(tag)

    if image_tags:
        text_content += "\n" + "\n".join(f"![Image]({tag})" for tag in image_tags)

    return {
        "text": text_content,
        "pages": [{"page_number": 1, "text": text_content}],
        "image_paths": image_paths,
        "pages_processed": len(doc.paragraphs),
    }


def parse_document(file_path: str, output_image_dir: str) -> Dict:
    """
    Auto-detect file type and parse.
    """
    if file_path.lower().endswith(".pdf"):
        return parse_pdf(file_path, output_image_dir)

    if file_path.lower().endswith(".docx"):
        return parse_docx(file_path, output_image_dir)

    raise ValueError("Unsupported file format")
